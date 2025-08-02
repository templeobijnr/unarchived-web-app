from django.http import JsonResponse
from rest_framework import viewsets, permissions, filters
from .models import DigitalProductGenome, DPGComponent, GeneratedAsset
from .serializers import DigitalProductGenomeSerializer, DPGComponentSerializer, GeneratedAssetSerializer, ApparelDPGExtensionSerializer
import tempfile
from google.cloud import vision
from google.cloud.vision_v1 import types
from google.oauth2 import service_account
import pytesseract
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from dpgs.models import DigitalProductGenome
from dpgs.serializers import DigitalProductGenomeSerializer
from agentcore.tools import dpg_builder_tool
from PIL import Image
from io import BytesIO
from files.models import UploadedFile
import fitz  # PyMuPDF
import docx
import csv
import pandas as pd
from rest_framework.decorators import action
from rest_framework.views import APIView
from django.shortcuts import render



class BearerTokenMixin:
    """
    Mixin to handle Bearer token authentication
    """
    def authenticate_bearer_token(self, request):
        """
        Extract and validate Bearer token from Authorization header
        Returns True if valid, False otherwise
        """
        auth_header = request.META.get('HTTP_AUTHORIZATION', '')
        
        if not auth_header.startswith('Bearer '):
            return False
            
        token = auth_header.split(' ')[1] if len(auth_header.split(' ')) > 1 else ''
        
        # Add your token validation logic here
        # For example, check against database, JWT validation, etc.
        
        # Simple example - replace with your actual validation
        VALID_TOKENS = [
            'your-secret-token-here',
            'test-token-123',
            # Add your valid tokens here
        ]
        
        return token in VALID_TOKENS
    
    def dispatch(self, request, *args, **kwargs):
        """
        Override dispatch to check authentication before processing
        """
        """if not self.authenticate_bearer_token(request):
            return JsonResponse({
                'error': 'Invalid or missing Bearer token'
            }, status=401)
        
        return super().dispatch(request, *args, **kwargs)"""

class UniversalGenerateDPG(APIView):
    parser_classes = [MultiPartParser, FormParser]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """Create a new Digital Product Genome based on the prompt and uploaded file."""
        prompt = request.data.get("prompt", "")
        file = request.data.get("file")

        extracted_text = ""

        if file:
            file_type = file.content_type
            file_name = file.name.lower()

            if file_type.startswith("image/"):
                image = Image.open(BytesIO(file.read()))
                extracted_text = pytesseract.image_to_string(image)
            elif file_name.endswith(".pdf"):
                pdf = fitz.open(stream=file.read(), filetype="pdf")
                for page in pdf:
                    extracted_text += page.get_text()
            elif file_name.endswith(".xls") or file_name.endswith(".xlsx"):
                try:
                    df = pd.read_excel(file, engine="openpyxl" if file_name.endswith(".xlsx") else "xlrd")
                    extracted_text = df.to_string(index=False)
                except Exception as e:
                    return Response({"error": f"Failed to read Excel file: {str(e)}"}, status=400)
            elif file_name.endswith(".docx"):
                doc = docx.Document(BytesIO(file.read()))
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            elif file_name.endswith(".txt") or file.content_type == "text/plain":
                extracted_text = file.read().decode("utf-8")
            elif file_name.endswith(".csv"):
                text = file.read().decode("utf-8")
                reader = csv.reader(text.splitlines())
                extracted_text = "\n".join(["\t".join(row) for row in reader])
            else:
                return Response({"error": "Unsupported file type."}, status=400)

        combined_input = f"{prompt}\n\n{extracted_text}".strip()

        try:
            dpg_data = dpg_builder_tool.invoke(combined_input)
            dpg = DigitalProductGenome.objects.create(
                title=dpg_data.get("title"),
                version=dpg_data.get("version", "1.0"),
                summary=dpg_data["data"].pop("summary", ""),
                data=dpg_data.get("data", {}),
                stage=dpg_data.get("stage", "created"),
                owner=request.user
            )

            # Initialize components for the DPG (if any)
            self._initialize_components(dpg)

            # Initialize apparel-specific extension (if applicable)
            self._initialize_apparel_extension(dpg)

            return Response(DigitalProductGenomeSerializer(dpg).data, status=201)
        except Exception as e:
            return Response({"error": str(e)}, status=500)
    def _initialize_components(self, dpg):
        """Initialize components for the DigitalProductGenome (DPG)."""
        components_data = dpg.data.get("components", [])
        for component_data in components_data:
            component_data['dpg'] = dpg.id
            component_serializer = DPGComponentSerializer(data=component_data)
            if component_serializer.is_valid():
                component_serializer.save()

    def _initialize_apparel_extension(self, dpg):
        """Initialize apparel-specific extension if needed."""
        if dpg.data.get("category") == "apparel":
            apparel_extension_data = {
                "dpg": dpg.id,
                "fabric_composition": dpg.data.get("fabric_composition", ""),
                "sizing_chart": dpg.data.get("sizing_chart", {}),
                "fit_information": dpg.data.get("fit_information", ""),
                "construction_details": dpg.data.get("construction_details", ""),
                "hardware_components": dpg.data.get("hardware_components", {}),
                "colorways": dpg.data.get("colorways", {}),
                "cost_breakdown": dpg.data.get("cost_breakdown", {}),
                "manufacturing_data": dpg.data.get("manufacturing_data", {})
            }
            apparel_extension_serializer = ApparelDPGExtensionSerializer(data=apparel_extension_data)
            if apparel_extension_serializer.is_valid():
                apparel_extension_serializer.save()
    
    def handle_file_upload(self, file):
        """Process and extract text from uploaded files."""
        extracted_text = ""
        file_type = file.content_type
        file_name = file.name.lower()

        if file_type.startswith("image/"):
            image = Image.open(BytesIO(file.read()))
            extracted_text = pytesseract.image_to_string(image)
        elif file_name.endswith(".pdf"):
            pdf = fitz.open(stream=file.read(), filetype="pdf")
            for page in pdf:
                extracted_text += page.get_text()
        elif file_name.endswith(".xls") or file_name.endswith(".xlsx"):
            df = pd.read_excel(file, engine="openpyxl" if file_name.endswith(".xlsx") else "xlrd")
            extracted_text = df.to_string(index=False)
        elif file_name.endswith(".docx"):
            doc = docx.Document(BytesIO(file.read()))
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_name.endswith(".txt") or file.content_type == "text/plain":
            extracted_text = file.read().decode("utf-8")
        elif file_name.endswith(".csv"):
            text = file.read().decode("utf-8")
            reader = csv.reader(text.splitlines())
            extracted_text = "\n".join(["\t".join(row) for row in reader])

        return extracted_text

class DigitalProductGenomeViewSet(viewsets.ModelViewSet):
    serializer_class = DigitalProductGenomeSerializer
    permission_classes = [permissions.IsAuthenticated]
    queryset = DigitalProductGenome.objects.all()  
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['title', 'description', 'version']
    ordering_fields = ['created_at', 'updated_at', 'version']

    def get_queryset(self):
        return DigitalProductGenome.objects.filter(owner=self.request.user)

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)

    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def add_component(self, request, pk=None):
        """Add a component to the DPG."""
        dpg = self.get_object()
        component_data = request.data
        component_data['dpg'] = dpg.id
        component_serializer = DPGComponentSerializer(data=component_data)
        if component_serializer.is_valid():
            component_serializer.save()
            return Response(component_serializer.data, status=status.HTTP_201_CREATED)
        return Response(component_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def add_asset(self, request, pk=None):
        """Add an asset (image, PDF, etc.) to the DPG."""
        dpg = self.get_object()
        asset_data = request.data
        asset_data['dpg'] = dpg.id
        asset_serializer = GeneratedAssetSerializer(data=asset_data)
        if asset_serializer.is_valid():
            asset_serializer.save()
            return Response(asset_serializer.data, status=status.HTTP_201_CREATED)
        return Response(asset_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def enhance_apparel(self, request, pk=None):
        """Add apparel specialization to the DPG."""
        dpg = self.get_object()
        apparel_data = request.data
        apparel_data['dpg'] = dpg.id
        apparel_extension_serializer = ApparelDPGExtensionSerializer(data=apparel_data)
        if apparel_extension_serializer.is_valid():
            apparel_extension_serializer.save()
            return Response(apparel_extension_serializer.data, status=status.HTTP_201_CREATED)
        return Response(apparel_extension_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def create_live_spec(self, request, pk=None):
        """Generate and create a shareable live spec for the DPG."""
        dpg = self.get_object()
        # Assuming the Live Spec is just a URL, but can be further enhanced.
        live_spec_link = f"/dpgs/{dpg.id}/view/"
        return Response({"live_spec_link": live_spec_link}, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['get'], permission_classes=[IsAuthenticated])
    def manufacturing_analysis(self, request, pk=None):
        """Get manufacturability insights based on DPG data."""
        dpg = self.get_object()
        # Call an external tool or function to fetch manufacturability insights.
        insights = manufacturability_insights_tool(dpg.data)  # Assumed tool for manufacturability analysis
        return Response({"insights": insights}, status=status.HTTP_200_OK)
    
class ExportDPG(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        try:
            dpg = DigitalProductGenome.objects.get(pk=pk, owner=request.user)
            return JsonResponse({
                "title": dpg.title,
                "version": dpg.version,
                "stage": dpg.stage,
                "data": dpg.data,
                "owner_id": dpg.owner.id,
                "created_at": dpg.created_at.isoformat(),
                "updated_at": dpg.updated_at.isoformat(),
            }, json_dumps_params={'indent': 4})
        except DigitalProductGenome.DoesNotExist:
            return JsonResponse({"error": "DPG not found"}, status=404)


def dpg_testing_ui(request):
    """Render the DPG testing UI template."""
    return render(request, 'dpg_testing_ui.html')